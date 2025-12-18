"""
Модуль для генерации отчетов в формате .docx.
Создает профессиональные отчеты по результатам экспериментов.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime
import os
import matplotlib.pyplot as plt


class ReportGenerator:
    """
    Класс для генерации отчетов по лабораторной работе.
    """
    
    def __init__(self):
        """Инициализация генератора отчетов."""
        self.doc = Document()
        self._setup_styles()
    
    def _setup_styles(self):
        """Настройка стилей документа."""
        # Настройка стиля Normal
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        
        # Создание стиля для заголовков
        try:
            heading_style = self.doc.styles.add_style('CustomHeading1', WD_STYLE_TYPE.PARAGRAPH)
            heading_font = heading_style.font
            heading_font.name = 'Times New Roman'
            heading_font.size = Pt(14)
            heading_font.bold = True
            heading_font.color.rgb = RGBColor(0, 0, 0)
        except:
            pass  # Стиль уже существует
    
    def add_title_page(self, university: str, department: str, title: str, 
                       authors: list, group: str, date: str):
        """
        Добавление титульного листа.
        
        Args:
            university: Название университета
            department: Название кафедры
            title: Название работы
            authors: Список авторов
            group: Группа
            date: Дата
        """
        # Университет (вверху, по центру)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(university)
        run.font.size = Pt(14)
        run.font.bold = True
        
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(department)
        run.font.size = Pt(12)
        
        # Пустые строки
        for _ in range(8):
            self.doc.add_paragraph()
        
        # Название работы
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.font.size = Pt(16)
        run.font.bold = True
        
        # Тип работы
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('Курсовая работа')
        run.font.size = Pt(14)
        
        # Пустые строки
        for _ in range(10):
            self.doc.add_paragraph()
        
        # Авторы (справа)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run('Выполнили:\n')
        run.font.size = Pt(12)
        for author in authors:
            run = p.add_run(f'{author}\n')
            run.font.size = Pt(12)
        run = p.add_run(f'Группа: {group}')
        run.font.size = Pt(12)
        
        # Дата (внизу, по центру)
        for _ in range(3):
            self.doc.add_paragraph()
        
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(date)
        run.font.size = Pt(12)
        
        self.doc.add_page_break()
    
    def add_heading(self, text: str, level: int = 1):
        """
        Добавление заголовка.
        
        Args:
            text: Текст заголовка
            level: Уровень заголовка (1-3)
        """
        heading = self.doc.add_heading(text, level=level)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Настройка шрифта
        for run in heading.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14 if level == 1 else 13 if level == 2 else 12)
            run.font.color.rgb = RGBColor(0, 0, 0)
    
    def add_paragraph(self, text: str, bold: bool = False, alignment: str = 'left'):
        """
        Добавление параграфа.
        
        Args:
            text: Текст параграфа
            bold: Жирный шрифт
            alignment: Выравнивание ('left', 'center', 'right', 'justify')
        """
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        if bold:
            run.font.bold = True
        
        alignment_map = {
            'left': WD_ALIGN_PARAGRAPH.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right': WD_ALIGN_PARAGRAPH.RIGHT,
            'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
        }
        p.alignment = alignment_map.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)
        
        return p
    
    def add_bullet_list(self, items: list):
        """
        Добавление маркированного списка.
        
        Args:
            items: Список элементов
        """
        for item in items:
            p = self.doc.add_paragraph(item, style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.5)
    
    def add_numbered_list(self, items: list):
        """
        Добавление нумерованного списка.
        
        Args:
            items: Список элементов
        """
        for item in items:
            p = self.doc.add_paragraph(item, style='List Number')
            p.paragraph_format.left_indent = Inches(0.5)
    
    def add_image(self, image_path: str, width: float = 5.0):
        """
        Добавление изображения.
        
        Args:
            image_path: Путь к изображению
            width: Ширина в дюймах
        """
        if os.path.exists(image_path):
            self.doc.add_picture(image_path, width=Inches(width))
            last_paragraph = self.doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def add_table(self, data: list, headers: list = None):
        """
        Добавление таблицы.
        
        Args:
            data: Данные таблицы (список списков)
            headers: Заголовки столбцов
        """
        rows = len(data) + (1 if headers else 0)
        cols = len(data[0]) if data else 0
        
        table = self.doc.add_table(rows=rows, cols=cols)
        table.style = 'Light Grid Accent 1'
        
        # Заголовки
        if headers:
            for i, header in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = header
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Данные
        start_row = 1 if headers else 0
        for i, row_data in enumerate(data):
            for j, cell_data in enumerate(row_data):
                table.rows[start_row + i].cells[j].text = str(cell_data)
    
    def generate_experiment_report(self, user_name: str, group: str, 
                                   calculator, results: dict, 
                                   diagram_path: str = None,
                                   stress_path: str = None):
        """
        Генерация полного отчета по эксперименту.
        
        Args:
            user_name: ФИО пользователя
            group: Группа
            calculator: Экземпляр TorsionCalculator
            results: Словарь с результатами расчетов
            diagram_path: Путь к диаграмме T-φ
            stress_path: Путь к графику распределения напряжений
        
        Returns:
            Путь к созданному файлу
        """
        # Титульный лист
        self.add_title_page(
            university='РОССИЙСКИЙ УНИВЕРСИТЕТ ДРУЖБЫ НАРОДОВ ИМЕНИ ПАТРИСА ЛУМУМБЫ',
            department='Кафедра сопротивления материалов',
            title='Определение модуля упругости второго рода\nпри кручении стали, чугуна, дерева',
            authors=[user_name],
            group=group,
            date=datetime.now().strftime('%d.%m.%Y')
        )
        
        # 1. Цель работы
        self.add_heading('1. Цель работы', level=1)
        self.add_paragraph(
            'Экспериментальное определение модуля сдвига G и механических '
            'характеристик материала при кручении: предела пропорциональности, '
            'предела текучести, предела прочности и максимального остаточного сдвига.',
            alignment='justify'
        )
        
        # 2. Исходные данные
        self.add_heading('2. Исходные данные', level=1)
        self.add_paragraph(f'Материал образца: {calculator.material}', bold=True)
        
        initial_data = [
            ['Параметр', 'Обозначение', 'Значение', 'Единица измерения'],
            ['Диаметр образца', 'D', f'{calculator.D * 1000:.2f}', 'мм'],
            ['Длина образца', 'L', f'{calculator.L * 1000:.2f}', 'мм'],
            ['Полярный момент инерции', 'Jp', f'{results["Jp"]:.6e}', 'м⁴'],
            ['Полярный момент сопротивления', 'Wp', f'{results["Wp"]:.6e}', 'м³'],
        ]
        self.add_table(initial_data, headers=initial_data[0])
        initial_data.pop(0)
        self.add_table(initial_data)
        
        # 3. Теоретические основы
        self.add_heading('3. Теоретические основы', level=1)
        self.add_paragraph(
            'При кручении валов круглого сечения справедливы следующие гипотезы:',
            alignment='justify'
        )
        self.add_bullet_list([
            'Плоские поперечные сечения остаются плоскими после деформации',
            'Радиусы сечений не искривляются',
            'Расстояния между сечениями не изменяются'
        ])
        
        self.add_paragraph('\nОсновные расчетные формулы:', bold=True)
        
        formulas = [
            '1. Относительный сдвиг: γ = (φ·D)/(2ℓ)',
            '2. Закон Гука при кручении: φ = T·ℓ/(G·Jp)',
            '3. Модуль сдвига: G = (T·ℓ)/(φ·Jp)',
            '4. Максимальное касательное напряжение: τmax = T/Wp',
            '5. Предел прочности при кручении: τB = Tk/Wp'
        ]
        self.add_numbered_list(formulas)
        
        # 4. Результаты эксперимента
        self.add_heading('4. Результаты эксперимента и расчетов', level=1)
        
        results_data = [
            ['Параметр', 'Значение', 'Единица измерения'],
            ['Модуль сдвига (эксп.)', f'{results["G_experimental"]:.2f}', 'МПа'],
            ['Модуль сдвига (эталон)', f'{results["G_reference"]:.2f}', 'МПа'],
            ['Относительная погрешность', f'{results["relative_error"]:.2f}', '%'],
            ['Максимальный момент', f'{results["T_max"]:.2f}', 'Н·м'],
            ['Угол при T_max', f'{results["phi_max"]:.5f}', 'рад'],
            ['Макс. касательное напряжение', f'{results["tau_max"]:.2f}', 'МПа'],
            ['Макс. остаточный сдвиг', f'{results["gamma_max"]:.5f}', 'рад'],
        ]
        self.add_table(results_data)
        
        # 5. Графики
        self.add_heading('5. Диаграммы и графики', level=1)
        
        if diagram_path and os.path.exists(diagram_path):
            self.add_paragraph('5.1. Диаграмма кручения T-φ:', bold=True)
            self.add_image(diagram_path, width=5.5)
            self.add_paragraph('')
        
        if stress_path and os.path.exists(stress_path):
            self.add_paragraph('5.2. Распределение касательных напряжений по сечению:', bold=True)
            self.add_image(stress_path, width=5.5)
        
        # 6. Выводы
        self.add_heading('6. Выводы', level=1)
        
        conclusions = [
            f'Экспериментально определен модуль сдвига материала "{calculator.material}": '
            f'G = {results["G_experimental"]:.2f} МПа.',
            
            f'Относительная погрешность по сравнению с эталонным значением составила '
            f'{results["relative_error"]:.2f}%.',
            
            f'Максимальное касательное напряжение в образце при T_max = {results["T_max"]:.2f} Н·м '
            f'составило τ_max = {results["tau_max"]:.2f} МПа.',
            
            'На диаграмме T-φ четко виден линейный участок, соответствующий упругой '
            'деформации материала, что подтверждает справедливость закона Гука при кручении.',
            
            f'Характер разрушения материала: {self._get_failure_description(calculator.material)}'
        ]
        
        for i, conclusion in enumerate(conclusions, 1):
            self.add_paragraph(f'{i}. {conclusion}', alignment='justify')
        
        # Сохранение
        filename = f'Отчет_{user_name.replace(" ", "_")}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
        self.doc.save(filename)
        return filename
    
    def _get_failure_description(self, material: str) -> str:
        """Получение описания характера разрушения."""
        descriptions = {
            'Сталь': 'разрушение по винтовой поверхности под углом 45° (срез по касательным напряжениям)',
            'Чугун': 'разрушение по плоскости, перпендикулярной оси (отрыв по нормальным напряжениям)',
            'Дерево': 'расслоение вдоль волокон'
        }
        return descriptions.get(material, 'типичное для данного материала')
    
    def save(self, filename: str):
        """
        Сохранение документа.
        
        Args:
            filename: Имя файла
        """
        self.doc.save(filename)
        return filename

