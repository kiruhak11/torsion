from docx import Document
from docx.shared import Inches

def generate_docx(material, L, diameter, moment, angle, G_basic, G_eff,
                  elastic_limit, failure_angle, graph_img_path, filename="report.docx",
                  experiment_results=None):
    """
    Генерирует DOCX-отчёт с исходными данными, рассчитанными значениями и графиком.
    Если передан experiment_results — вставляет параметры эксперимента.
    """
    doc = Document()
    doc.add_heading("Отчёт по экспериментальному определению прочности при кручении", 0)

    doc.add_heading("Исходные данные", level=1)
    doc.add_paragraph(f"Материал: {material}")
    doc.add_paragraph(f"Длина образца (мм): {L}")
    doc.add_paragraph(f"Диаметр образца (мм): {diameter}")
    doc.add_paragraph(f"Крутящий момент (Н·мм): {moment}")
    doc.add_paragraph(f"Введённый угол (°): {angle}")
    doc.add_paragraph(f"Предел эластичности (°): {elastic_limit}")
    doc.add_paragraph(f"Угол разрушения (°): {failure_angle}")

    doc.add_heading("Результаты расчёта", level=1)
    doc.add_paragraph(f"Базовый модуль упругости (G): {G_basic} МПа")
    doc.add_paragraph(f"Эффективный модуль упругости (G_eff): {G_eff} МПа")

    if experiment_results:
        doc.add_heading("Результаты по экспериментальным данным", level=1)
        doc.add_paragraph(f"Модуль G (линейный участок): {experiment_results.get('G_linear')} МПа")
        doc.add_paragraph(f"Предел пропорциональности (τпц): {experiment_results.get('τ_pcz')} МПа")
        doc.add_paragraph(f"Предел текучести по 0.3% сдвигу: {experiment_results.get('τ_0_3')} МПа")
        doc.add_paragraph(f"Предел прочности (τmax): {experiment_results.get('τ_max')} МПа")
        doc.add_paragraph(f"Максимальный относительный сдвиг (γmax): {experiment_results.get('γ_max')}")

    doc.add_heading("График скручивания", level=1)
    try:
        doc.add_picture(graph_img_path, width=Inches(6))
    except Exception:
        doc.add_paragraph("График не удалось вставить.")

    doc.add_heading("Пояснение к расчётам", level=1)
    doc.add_paragraph(
        "1. Модуль сдвига (второго рода) рассчитывается по формуле:\n"
        "   G = (T * L) / (J * φ), где J = (π * d⁴) / 32, φ — в радианах.\n\n"
        "2. В пределах упругости G_eff = k * G.\n"
        "3. В пластической области G_eff снижается линейно до 0.\n"
        "4. При анализе эксперимента определяется модуль упругости на линейном участке,\n"
        "   предел пропорциональности, предел текучести по 0.3% сдвигу и максимальный сдвиг."
    )

    doc.save(filename)
